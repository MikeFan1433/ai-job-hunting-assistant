"""Resume Export Module - Export resume to PDF and Word formats."""
import io
import os
from pathlib import Path
from typing import Dict, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class ResumeExporter:
    """Export resume to PDF and Word formats."""
    
    def __init__(self):
        """Initialize the resume exporter."""
        self.supported_formats = []
        if PDF_AVAILABLE:
            self.supported_formats.append("pdf")
        if DOCX_AVAILABLE:
            self.supported_formats.append("docx")

    def _build_plain_text_pdf_story(self, resume_text: str, title: str):
        """Build reportlab story for plain multi-line text (resume or arbitrary export)."""
        story = []
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#000000',
            spaceAfter=12,
            alignment=TA_CENTER
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2*inch))

        lines = resume_text.split('\n')
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.1*inch))
                continue

            if self._is_section_header(line):
                if current_section:
                    story.append(Spacer(1, 0.15*inch))

                section_style = ParagraphStyle(
                    'SectionHeader',
                    parent=styles['Heading2'],
                    fontSize=14,
                    textColor='#000000',
                    spaceAfter=6,
                    spaceBefore=12,
                    fontName='Helvetica-Bold'
                )
                story.append(Paragraph(line.upper(), section_style))
                current_section = line
            else:
                if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                    content = line.lstrip('•-*').strip()
                    para_style = ParagraphStyle(
                        'BulletPoint',
                        parent=styles['Normal'],
                        fontSize=10,
                        leftIndent=0.25*inch,
                        spaceAfter=4
                    )
                    story.append(Paragraph(f"• {content}", para_style))
                else:
                    if '|' in line or 'at' in line.lower():
                        para_style = ParagraphStyle(
                            'JobTitle',
                            parent=styles['Normal'],
                            fontSize=11,
                            fontName='Helvetica-Bold',
                            spaceAfter=4
                        )
                    else:
                        para_style = styles['Normal']

                    story.append(Paragraph(line, para_style))
        return story

    def export_plain_text_pdf_bytes(self, resume_text: str, title: str = "Document") -> Dict:
        """Render plain text to PDF in memory (no temp file on disk)."""
        if not PDF_AVAILABLE:
            return {
                "error": "PDF export not available. Please install reportlab: pip install reportlab"
            }
        try:
            buf = io.BytesIO()
            doc = SimpleDocTemplate(
                buf,
                pagesize=letter,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )
            doc.build(self._build_plain_text_pdf_story(resume_text, title))
            data = buf.getvalue()
            return {
                "status": "success",
                "data": data,
                "format": "pdf",
                "size_kb": round(len(data) / 1024, 2),
            }
        except Exception as e:
            return {"error": f"Failed to export PDF: {str(e)}"}
    
    def export_to_pdf(
        self,
        resume_text: str,
        output_path: str,
        title: str = "Resume"
    ) -> Dict:
        """
        Export resume to PDF format.
        
        Args:
            resume_text: Resume text content
            output_path: Path to save the PDF file
            title: Title for the document
        
        Returns:
            Dictionary with export status
        """
        if not PDF_AVAILABLE:
            return {
                "error": "PDF export not available. Please install reportlab: pip install reportlab"
            }
        
        try:
            path = Path(output_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            
            doc = SimpleDocTemplate(
                str(path),
                pagesize=letter,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )
            doc.build(self._build_plain_text_pdf_story(resume_text, title))
            
            return {
                "status": "success",
                "filepath": str(path),
                "format": "pdf",
                "size_kb": round(path.stat().st_size / 1024, 2)
            }
        
        except Exception as e:
            return {
                "error": f"Failed to export PDF: {str(e)}"
            }

    def _make_docx_from_plain_text(self, resume_text: str, title: str) -> "Document":
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        lines = resume_text.split('\n')
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            if self._is_section_header(line):
                if current_section:
                    doc.add_paragraph()
                doc.add_heading(line, level=2)
                current_section = line
            else:
                para = doc.add_paragraph()
                if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                    content = line.lstrip('•-*').strip()
                    para.style = 'List Bullet'
                    para.add_run(content)
                else:
                    if '|' in line or 'at' in line.lower():
                        run = para.add_run(line)
                        run.bold = True
                        run.font.size = Pt(11)
                    else:
                        para.add_run(line)
        return doc

    def export_plain_text_docx_bytes(self, resume_text: str, title: str = "Resume") -> Dict:
        """Build DOCX in memory (no temp file)."""
        if not DOCX_AVAILABLE:
            return {
                "error": "DOCX export not available. Please install python-docx: pip install python-docx"
            }
        try:
            doc = self._make_docx_from_plain_text(resume_text, title)
            buf = io.BytesIO()
            doc.save(buf)
            data = buf.getvalue()
            return {
                "status": "success",
                "data": data,
                "format": "docx",
                "size_kb": round(len(data) / 1024, 2),
            }
        except Exception as e:
            return {"error": f"Failed to export DOCX: {str(e)}"}
    
    def export_to_docx(
        self,
        resume_text: str,
        output_path: str,
        title: str = "Resume"
    ) -> Dict:
        """
        Export resume to Word (DOCX) format.
        
        Args:
            resume_text: Resume text content
            output_path: Path to save the DOCX file
            title: Title for the document
        
        Returns:
            Dictionary with export status
        """
        if not DOCX_AVAILABLE:
            return {
                "error": "DOCX export not available. Please install python-docx: pip install python-docx"
            }
        
        try:
            path = Path(output_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            doc = self._make_docx_from_plain_text(resume_text, title)
            doc.save(str(path))
            
            return {
                "status": "success",
                "filepath": str(path),
                "format": "docx",
                "size_kb": round(path.stat().st_size / 1024, 2)
            }
        
        except Exception as e:
            return {
                "error": f"Failed to export DOCX: {str(e)}"
            }
    
    def _is_section_header(self, line: str) -> bool:
        """Check if a line is a section header."""
        # Common section headers
        section_headers = [
            "work experience", "experience", "employment", "work history",
            "education", "educational background", "academic background",
            "skills", "technical skills", "professional skills",
            "projects", "project experience",
            "certifications", "awards", "honors",
            "languages", "interests", "references"
        ]
        
        line_lower = line.lower()
        
        # Check if it's all caps (likely a header)
        if line.isupper() and len(line) > 2:
            return True
        
        # Check if it matches common headers
        for header in section_headers:
            if header in line_lower and len(line) < 50:
                return True
        
        return False
    
    def export(
        self,
        resume_text: str,
        output_path: str,
        format: str = "pdf",
        title: str = "Resume"
    ) -> Dict:
        """
        Export resume to specified format.
        
        Args:
            resume_text: Resume text content
            output_path: Path to save the file
            format: Export format ("pdf" or "docx")
            title: Title for the document
        
        Returns:
            Dictionary with export status
        """
        if format.lower() == "pdf":
            return self.export_to_pdf(resume_text, output_path, title)
        elif format.lower() == "docx":
            return self.export_to_docx(resume_text, output_path, title)
        else:
            return {
                "error": f"Unsupported format: {format}. Supported formats: {', '.join(self.supported_formats)}"
            }
